import logging
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Optional, Dict, Any
from io import BytesIO

import fitz  # PyMuPDF
from docx import Document as DocxDocument
import markdown
from bs4 import BeautifulSoup

# OCR imports (optional - gracefully degrade if not available)
try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

from ..models.documents import DocumentType

logger = logging.getLogger(__name__)


class BaseParser(ABC):
    
    @abstractmethod
    def parse(self, file_path: Path) -> str:
        pass
    
    @abstractmethod
    def parse_bytes(self, content: bytes, filename: str) -> str:
        pass
    
    def extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        return {
            "filename": file_path.name,
            "size_bytes": file_path.stat().st_size if file_path.exists() else 0,
        }


class PDFParser(BaseParser):
    
    def parse(self, file_path: Path) -> str:
        logger.info(f"Parsing PDF: {file_path}")
        
        try:
            doc = fitz.open(file_path)
            
            text_parts = []
            for page_num, page in enumerate(doc):
                page_text = page.get_text("text")
                
                if page_text.strip():
                    text_parts.append(page_text)
                    
            doc.close()
            
            full_text = "\n\n".join(text_parts)
            
            logger.info(f"Extracted {len(full_text)} characters from {len(text_parts)} pages")
            return full_text
            
        except Exception as e:
            logger.error(f"Failed to parse PDF {file_path}: {e}")
            raise ValueError(f"Failed to parse PDF: {e}")
    
    def parse_bytes(self, content: bytes, filename: str) -> str:
        logger.info(f"Parsing PDF from bytes: {filename}")
        
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            
            text_parts = []
            for page in doc:
                page_text = page.get_text("text")
                if page_text.strip():
                    text_parts.append(page_text)
                    
            doc.close()
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Failed to parse PDF bytes {filename}: {e}")
            raise ValueError(f"Failed to parse PDF: {e}")
    
    def extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        base_meta = super().extract_metadata(file_path)
        
        try:
            doc = fitz.open(file_path)
            meta = doc.metadata
            base_meta.update({
                "page_count": doc.page_count,
                "title": meta.get("title", ""),
                "author": meta.get("author", ""),
                "subject": meta.get("subject", ""),
                "creator": meta.get("creator", ""),
                "creation_date": meta.get("creationDate", ""),
            })
            
            doc.close()
            
        except Exception as e:
            logger.warning(f"Could not extract PDF metadata: {e}")
            
        return base_meta


class DOCXParser(BaseParser):
    
    def parse(self, file_path: Path) -> str:
        logger.info(f"Parsing DOCX: {file_path}")
        
        try:
            doc = DocxDocument(file_path)
            
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            full_text = "\n\n".join(paragraphs)
            
            logger.info(f"Extracted {len(full_text)} characters from DOCX")
            return full_text
            
        except Exception as e:
            logger.error(f"Failed to parse DOCX {file_path}: {e}")
            raise ValueError(f"Failed to parse DOCX: {e}")
    
    def parse_bytes(self, content: bytes, filename: str) -> str:
        logger.info(f"Parsing DOCX from bytes: {filename}")
        
        try:
            doc = DocxDocument(BytesIO(content))
            
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            return "\n\n".join(paragraphs)
            
        except Exception as e:
            logger.error(f"Failed to parse DOCX bytes {filename}: {e}")
            raise ValueError(f"Failed to parse DOCX: {e}")


class MarkdownParser(BaseParser):
    
    def __init__(self):
        self.md = markdown.Markdown(
            extensions=['tables', 'fenced_code', 'nl2br']
        )
    
    def parse(self, file_path: Path) -> str:
        logger.info(f"Parsing Markdown: {file_path}")
        
        try:
            content = file_path.read_text(encoding='utf-8')
            return self._convert_to_text(content)
            
        except Exception as e:
            logger.error(f"Failed to parse Markdown {file_path}: {e}")
            raise ValueError(f"Failed to parse Markdown: {e}")
    
    def parse_bytes(self, content: bytes, filename: str) -> str:
        logger.info(f"Parsing Markdown from bytes: {filename}")
        
        try:
            text = content.decode('utf-8')
            return self._convert_to_text(text)
            
        except Exception as e:
            logger.error(f"Failed to parse Markdown bytes {filename}: {e}")
            raise ValueError(f"Failed to parse Markdown: {e}")
    
    def _convert_to_text(self, md_content: str) -> str:
        self.md.reset()
        html = self.md.convert(md_content)
        soup = BeautifulSoup(html, 'html.parser')
        text = soup.get_text(separator='\n')
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(line for line in lines if line)
        
        logger.info(f"Converted {len(md_content)} chars MD to {len(text)} chars text")
        return text


class TextParser(BaseParser):
    
    def parse(self, file_path: Path) -> str:
        logger.info(f"Parsing text file: {file_path}")
        
        try:
            try:
                content = file_path.read_text(encoding='utf-8')
            except UnicodeDecodeError:
                content = file_path.read_text(encoding='latin-1')
            
            content = content.replace('\r\n', '\n').replace('\r', '\n')
            
            if content.startswith('\ufeff'):
                content = content[1:]
            
            logger.info(f"Read {len(content)} characters from text file")
            return content
            
        except Exception as e:
            logger.error(f"Failed to parse text file {file_path}: {e}")
            raise ValueError(f"Failed to parse text file: {e}")
    
    def parse_bytes(self, content: bytes, filename: str) -> str:
        logger.info(f"Parsing text from bytes: {filename}")
        
        try:
            try:
                text = content.decode('utf-8')
            except UnicodeDecodeError:
                text = content.decode('latin-1')
            
            text = text.replace('\r\n', '\n').replace('\r', '\n')
            
            if text.startswith('\ufeff'):
                text = text[1:]
            
            return text
            
        except Exception as e:
            logger.error(f"Failed to parse text bytes {filename}: {e}")
            raise ValueError(f"Failed to parse text: {e}")


class ImageParser(BaseParser):
    """Parser for images using OCR (Optical Character Recognition)."""
    
    def __init__(self):
        if not OCR_AVAILABLE:
            logger.warning("OCR dependencies not available. Install with: pip install pytesseract Pillow")
    
    def parse(self, file_path: Path) -> str:
        logger.info(f"Parsing image with OCR: {file_path}")
        
        if not OCR_AVAILABLE:
            raise ValueError("OCR is not available. Please install pytesseract and Pillow.")
        
        try:
            image = Image.open(file_path)
            # Pre-process image for better OCR results
            image = self._preprocess_image(image)
            text = pytesseract.image_to_string(image)
            
            # Clean up the text
            text = self._clean_ocr_text(text)
            
            logger.info(f"Extracted {len(text)} characters from image via OCR")
            return text
            
        except Exception as e:
            logger.error(f"Failed to parse image {file_path}: {e}")
            raise ValueError(f"Failed to parse image with OCR: {e}")
    
    def parse_bytes(self, content: bytes, filename: str) -> str:
        logger.info(f"Parsing image from bytes with OCR: {filename}")
        
        if not OCR_AVAILABLE:
            raise ValueError("OCR is not available. Please install pytesseract and Pillow.")
        
        try:
            image = Image.open(BytesIO(content))
            # Pre-process image for better OCR results
            image = self._preprocess_image(image)
            text = pytesseract.image_to_string(image)
            
            # Clean up the text
            text = self._clean_ocr_text(text)
            
            logger.info(f"Extracted {len(text)} characters from image via OCR")
            return text
            
        except Exception as e:
            logger.error(f"Failed to parse image bytes {filename}: {e}")
            raise ValueError(f"Failed to parse image with OCR: {e}")
    
    def _preprocess_image(self, image: 'Image.Image') -> 'Image.Image':
        """Pre-process image for better OCR accuracy."""
        # Convert to RGB if necessary (handles PNG with alpha, etc.)
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Convert to grayscale for better OCR
        image = image.convert('L')
        
        # Optional: Increase contrast and size for small images
        width, height = image.size
        if width < 1000 or height < 1000:
            # Scale up small images
            scale = max(1000 / width, 1000 / height, 1)
            if scale > 1:
                new_size = (int(width * scale), int(height * scale))
                image = image.resize(new_size, Image.Resampling.LANCZOS)
        
        return image
    
    def _clean_ocr_text(self, text: str) -> str:
        """Clean up OCR output text."""
        # Remove excessive whitespace
        lines = [line.strip() for line in text.split('\n')]
        # Remove empty lines but keep paragraph breaks
        cleaned_lines = []
        prev_empty = False
        for line in lines:
            if line:
                cleaned_lines.append(line)
                prev_empty = False
            elif not prev_empty:
                cleaned_lines.append('')
                prev_empty = True
        
        return '\n'.join(cleaned_lines).strip()
    
    def extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        base_meta = super().extract_metadata(file_path)
        
        if OCR_AVAILABLE:
            try:
                image = Image.open(file_path)
                base_meta.update({
                    "width": image.width,
                    "height": image.height,
                    "format": image.format,
                    "mode": image.mode,
                })
            except Exception as e:
                logger.warning(f"Could not extract image metadata: {e}")
        
        return base_meta


PARSER_REGISTRY: Dict[DocumentType, type] = {
    DocumentType.PDF: PDFParser,
    DocumentType.DOCX: DOCXParser,
    DocumentType.MARKDOWN: MarkdownParser,
    DocumentType.TXT: TextParser,
    DocumentType.IMAGE: ImageParser,
}

_parser_instances: Dict[DocumentType, BaseParser] = {}


def get_parser(doc_type: DocumentType) -> BaseParser:
    if doc_type not in _parser_instances:
        parser_class = PARSER_REGISTRY.get(doc_type)
        if parser_class is None:
            raise ValueError(f"No parser available for document type: {doc_type}")
        _parser_instances[doc_type] = parser_class()
    
    return _parser_instances[doc_type]


def detect_document_type(filename: str) -> DocumentType:
    extension = Path(filename).suffix.lower()
    
    extension_map = {
        '.pdf': DocumentType.PDF,
        '.docx': DocumentType.DOCX,
        '.doc': DocumentType.DOCX,  # Try DOCX parser for .doc too
        '.md': DocumentType.MARKDOWN,
        '.markdown': DocumentType.MARKDOWN,
        '.txt': DocumentType.TXT,
        '.text': DocumentType.TXT,
        '.log': DocumentType.TXT,
        '.csv': DocumentType.TXT,  # Treat CSV as text for now
        '.json': DocumentType.TXT,  # Treat JSON as text
        '.xml': DocumentType.TXT,   # Treat XML as text
        '.html': DocumentType.TXT,  # Could add HTML parser later
        '.htm': DocumentType.TXT,
        # Image types (OCR)
        '.png': DocumentType.IMAGE,
        '.jpg': DocumentType.IMAGE,
        '.jpeg': DocumentType.IMAGE,
        '.gif': DocumentType.IMAGE,
        '.bmp': DocumentType.IMAGE,
        '.tiff': DocumentType.IMAGE,
        '.tif': DocumentType.IMAGE,
        '.webp': DocumentType.IMAGE,
    }
    
    doc_type = extension_map.get(extension)
    if doc_type is None:
        raise ValueError(f"Unsupported file type: {extension}")
    
    return doc_type


def parse_document(file_path: Path) -> tuple[str, DocumentType]:
    doc_type = detect_document_type(file_path.name)
    parser = get_parser(doc_type)
    text = parser.parse(file_path)
    return text, doc_type


def parse_document_bytes(content: bytes, filename: str) -> tuple[str, DocumentType]:
    doc_type = detect_document_type(filename)
    parser = get_parser(doc_type)
    text = parser.parse_bytes(content, filename)
    return text, doc_type

